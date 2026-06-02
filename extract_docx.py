from __future__ import annotations

import json
from pathlib import Path

from docx import Document


def extract(docx_path: Path) -> dict:
    doc = Document(str(docx_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    tables: list[list[list[str]]] = []
    for t in doc.tables:
        rows: list[list[str]] = []
        for r in t.rows:
            rows.append([c.text.strip() for c in r.cells])
        tables.append(rows)

    return {"paragraphs": paragraphs, "tables": tables}


def main() -> None:
    docx_path = Path(__file__).with_name("养兔管理系统完整技术文档.docx")
    out_json = Path(__file__).with_name("养兔管理系统完整技术文档.extracted.json")
    out_md = Path(__file__).with_name("养兔管理系统完整技术文档.extracted.md")

    data = extract(docx_path)
    out_json.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    md_lines: list[str] = []
    md_lines.append(f"# {docx_path.name}（提取版）")
    md_lines.append("")
    md_lines.append("## 段落")
    md_lines.append("")
    for i, p in enumerate(data["paragraphs"], start=1):
        md_lines.append(f"{i}. {p}")

    md_lines.append("")
    md_lines.append("## 表格")
    md_lines.append("")
    for ti, table in enumerate(data["tables"], start=1):
        md_lines.append(f"### 表格 {ti}")
        md_lines.append("")
        for ri, row in enumerate(table, start=1):
            md_lines.append(f"- 行 {ri}: " + " | ".join(row))
        md_lines.append("")

    out_md.write_text("\n".join(md_lines), encoding="utf-8")
    print(f"Wrote: {out_json}")
    print(f"Wrote: {out_md}")


if __name__ == "__main__":
    main()
